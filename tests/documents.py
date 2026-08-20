"""Builders for real PDF and DOCX bytes used across the extractor tests.

Generating genuine documents (rather than stubbing the libraries) means the
extractor tests exercise the same code path as production.
"""

from __future__ import annotations

from io import BytesIO

import pypdf
from docx import Document


def build_pdf(lines: list[str], *, pages: int = 1) -> bytes:
    """Return a minimal but valid PDF containing the given lines of text."""
    objects: list[bytes] = []
    page_object_ids = [3 + index * 2 for index in range(pages)]

    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = b" ".join(f"{object_id} 0 R".encode() for object_id in page_object_ids)
    objects.append(b"<< /Type /Pages /Kids [" + kids + b"] /Count " + str(pages).encode() + b" >>")

    font_id = 3 + pages * 2
    for index, page_object_id in enumerate(page_object_ids):
        content_id = page_object_id + 1
        page_lines = lines if pages == 1 else [f"page {index + 1}: {line}" for line in lines]
        stream = _text_stream(page_lines)
        objects.append(
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 " + str(font_id).encode() + b" 0 R >> >> "
            b"/Contents " + str(content_id).encode() + b" 0 R >>"
        )
        objects.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
        )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    return _assemble_pdf(objects)


def build_encrypted_pdf(password: str = "secret") -> bytes:
    """Return a password-protected PDF, which no extractor may open."""
    reader = pypdf.PdfReader(BytesIO(build_pdf(["confidential"])))
    writer = pypdf.PdfWriter()
    writer.add_page(reader.pages[0])
    writer.encrypt(password)

    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def build_docx(paragraphs: list[str], *, table_rows: list[list[str]] | None = None) -> bytes:
    """Return a real .docx file with the given paragraphs and optional table."""
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for cell_index, value in enumerate(row):
                table.cell(row_index, cell_index).text = value

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _text_stream(lines: list[str]) -> bytes:
    operations = ["BT", "/F1 12 Tf", "72 720 Td", "14 TL"]
    for index, line in enumerate(lines):
        escaped = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        operations.append(f"({escaped}) Tj" if index == 0 else f"T* ({escaped}) Tj")
    operations.append("ET")
    return "\n".join(operations).encode("latin-1")


def _assemble_pdf(objects: list[bytes]) -> bytes:
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += str(number).encode() + b" 0 obj\n" + body + b"\nendobj\n"

    xref_offset = len(out)
    size = len(objects) + 1
    out += b"xref\n0 " + str(size).encode() + b"\n0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        b"trailer\n<< /Size " + str(size).encode() + b" /Root 1 0 R >>\n"
        b"startxref\n" + str(xref_offset).encode() + b"\n%%EOF\n"
    )
    return bytes(out)
