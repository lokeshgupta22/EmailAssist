"""Stage 3a - pull plain text out of a document.

These functions do the actual reading and are deliberately dependency-heavy:
``pypdf`` and ``python-docx`` are large parsers that will happily crash, hang
or allocate wildly when fed a malformed file. That is why nothing here is
called directly by the application - :mod:`app.pipeline.sandbox` runs it in a
separate, resource-limited process.

Adding a new document type means adding one function and one registry entry;
no caller changes.
"""

from __future__ import annotations

from collections.abc import Callable
from io import BytesIO

import pypdf
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.config import DOCX_MIME, PDF_MIME
from app.pipeline.text import normalise_whitespace

TRUNCATION_MARKER = " [truncated]"


class ExtractionError(RuntimeError):
    """Raised when a document cannot be read. Always carries a printable reason."""


def extract_text(mime: str, payload: bytes, *, max_chars: int) -> str:
    """Extract plain text from a supported document.

    Raises :class:`ExtractionError` for anything that cannot be read, so the
    caller never has to reason about library-specific exceptions.
    """
    if not payload:
        raise ExtractionError("the document is empty")

    extractor = _EXTRACTORS.get(mime)
    if extractor is None:
        raise ExtractionError(f"{mime} documents are not supported")

    text = normalise_whitespace(extractor(payload))
    return _truncate(text, max_chars)


def supported_mime_types() -> frozenset[str]:
    """The document types this module can read."""
    return frozenset(_EXTRACTORS)


def _extract_pdf(payload: bytes) -> str:
    try:
        reader = pypdf.PdfReader(BytesIO(payload))
        if reader.is_encrypted:
            raise ExtractionError("the PDF is password protected, so it was not opened")
        pages = [page.extract_text() or "" for page in reader.pages]
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"the PDF could not be read: {_short(exc)}") from exc

    return "\n\n".join(page for page in pages if page.strip())


def _extract_docx(payload: bytes) -> str:
    try:
        document = Document(BytesIO(payload))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    except PackageNotFoundError as exc:
        raise ExtractionError("the file is not a readable Word document") from exc
    except Exception as exc:
        raise ExtractionError(f"the document could not be read: {_short(exc)}") from exc

    return "\n".join(block.strip() for block in blocks if block.strip())


def _truncate(text: str, max_chars: int) -> str:
    """Cap the text so one huge attachment cannot dominate the model's context."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + TRUNCATION_MARKER


def _short(exc: Exception) -> str:
    """Keep library error messages to one short, printable line."""
    message = str(exc).strip().splitlines()
    return (message[0][:120] if message else exc.__class__.__name__) or exc.__class__.__name__


_EXTRACTORS: dict[str, Callable[[bytes], str]] = {
    PDF_MIME: _extract_pdf,
    DOCX_MIME: _extract_docx,
}
